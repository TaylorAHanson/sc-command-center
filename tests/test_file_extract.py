"""Standalone tests for upload parsing in `file_extract`.

Fixtures are built in memory so this runs without credentials or sample files.
The PDF case needs a PDF writer, which is not a pinned dependency, so it skips
itself rather than failing when matplotlib is absent.
"""
import io
import json
import os
import sys
import tempfile

# matplotlib writes a cache on first import and hard-fails when HOME is not
# writable, which is the case in some sandboxes.
os.environ.setdefault("MPLCONFIGDIR", tempfile.gettempdir())

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "server"))

from services.file_extract import (  # noqa: E402
    KIND_DATA,
    KIND_DOCUMENT,
    KIND_IMAGE,
    KIND_TABLE,
    Extracted,
    UnsupportedFile,
    extract,
    file_card,
    load_table,
    sheet_names,
    sniff_kind,
    supported_extensions,
)


def _workbook_bytes(sheets):
    import pandas as pd

    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for name, rows in sheets.items():
            pd.DataFrame(rows).to_excel(writer, sheet_name=name, index=False)
    return buffer.getvalue()


def _docx_bytes():
    import docx

    document = docx.Document()
    document.add_heading("Quarterly Review", level=1)
    document.add_paragraph("Revenue grew in every region.")
    document.add_heading("Risks", level=2)
    document.add_paragraph("Supplier concentration remains high.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Region"
    table.cell(0, 1).text = "Revenue"
    table.cell(1, 0).text = "EMEA"
    table.cell(1, 1).text = "1200"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_bytes():
    """A two-page PDF, or None when no writer is available."""
    try:
        import matplotlib

        matplotlib.use("pdf")
        from matplotlib.backends.backend_pdf import PdfPages
        from matplotlib.figure import Figure
    except Exception:
        return None

    buffer = io.BytesIO()
    with PdfPages(buffer) as pdf:
        for text in ("Alpha page content", "Bravo page content"):
            figure = Figure(figsize=(4, 3))
            figure.text(0.1, 0.5, text)
            pdf.savefig(figure)
    return buffer.getvalue()


def test_supported_extensions_map_to_the_four_kinds():
    mapping = supported_extensions()
    assert mapping[".xlsx"] == KIND_TABLE
    assert mapping[".pdf"] == KIND_DOCUMENT
    assert mapping[".png"] == KIND_IMAGE
    assert mapping[".json"] == KIND_DATA
    # A caller mutating the advertised map must not reconfigure the module.
    mapping[".exe"] = KIND_TABLE
    assert ".exe" not in supported_extensions()


def test_sniff_kind_uses_extension_then_mime():
    assert sniff_kind("q1.CSV") == KIND_TABLE
    assert sniff_kind("notes", "text/markdown") == KIND_DOCUMENT
    assert sniff_kind("shot", "image/avif") == KIND_IMAGE


def test_sniff_kind_rejects_an_unsupported_extension():
    for name in ("installer.exe", "archive.zip", "clip.mp4"):
        try:
            sniff_kind(name)
        except UnsupportedFile as exc:
            assert "Supported" in str(exc), exc
        else:
            raise AssertionError(f"{name} should not be supported")


def test_multi_sheet_workbook_round_trips_by_sheet_name():
    data = _workbook_bytes(
        {
            "Summary": [{"metric": "revenue", "value": 10}],
            "Line Items!": [{"sku": f"S{i}", "qty": i} for i in range(5)],
        }
    )
    result = extract("book.xlsx", data)

    assert result.kind == KIND_TABLE
    assert result.tables is not None
    assert sheet_names(result.tables) == ["Summary", "Line Items!"]
    # The primary sheet is the one with the most rows, regardless of position.
    assert result.profile["primary_sheet"] == "Line Items!"
    assert load_table(result.tables).shape == (5, 2)

    summary = load_table(result.tables, "Summary")
    assert list(summary.columns) == ["metric", "value"]
    assert summary.iloc[0]["value"] == 10

    members = {sheet["name"]: sheet["member"] for sheet in result.profile["sheets"]}
    assert members["Line Items!"] == "Line_Items.parquet", members
    columns = {c["name"]: c for c in result.profile["sheets"][1]["columns"]}
    assert columns["qty"]["unique"] == 5
    assert columns["sku"]["sample_values"][:2] == ["S0", "S1"]

    try:
        load_table(result.tables, "Nope")
    except KeyError as exc:
        assert "Available" in str(exc)
    else:
        raise AssertionError("an unknown sheet name should raise")


def test_csv_with_a_non_utf8_encoding_and_semicolon_delimiter():
    rows = ["name;city"] + [f"Müller {i};Köln" for i in range(20)]
    data = "\n".join(rows).encode("cp1252")
    assert b"\xfc" in data

    result = extract("people.csv", data)
    frame = load_table(result.tables)
    assert list(frame.columns) == ["name", "city"], frame.columns
    assert frame.iloc[0]["name"] == "Müller 0", frame.iloc[0]["name"]
    assert frame.iloc[0]["city"] == "Köln"
    assert result.profile["sheets"][0]["name"] == "data"
    assert result.profile["sheets"][0]["truncated"] is False


def test_tsv_yields_a_single_sheet_named_data():
    result = extract("t.tsv", b"a\tb\n1\t2\n3\t4\n")
    assert sheet_names(result.tables) == ["data"]
    assert load_table(result.tables).shape == (2, 2)


def test_json_array_of_objects_becomes_a_table():
    payload = json.dumps([{"id": 1, "name": "a"}, {"id": 2, "name": "b"}]).encode()
    result = extract("rows.json", payload)

    assert result.kind == KIND_TABLE
    assert sheet_names(result.tables) == ["data"]
    frame = load_table(result.tables)
    assert list(frame["name"]) == ["a", "b"]


def test_ndjson_of_objects_becomes_a_table():
    payload = b'{"id": 1}\n{"id": 2}\n{"id": 3}\n'
    result = extract("rows.ndjson", payload)
    assert result.kind == KIND_TABLE
    assert load_table(result.tables).shape == (3, 1)


def test_json_object_stays_data_with_keys_and_depth():
    payload = json.dumps({"config": {"nested": {"deep": [1, 2]}}, "version": 3}).encode()
    result = extract("config.json", payload)

    assert result.kind == KIND_DATA
    assert result.tables is None
    assert result.text == ""
    assert result.profile["json_kind"] == "object"
    assert result.profile["keys"] == ["config", "version"]
    assert result.profile["depth"] == 4, result.profile["depth"]
    assert '"version": 3' in result.profile["preview"]
    assert len(result.profile["preview"]) <= 2000


def test_invalid_json_warns_instead_of_raising():
    result = extract("broken.json", b'{"a": 1,,,}')
    assert result.kind == KIND_DATA
    assert result.profile["json_kind"] == "invalid"
    assert result.warnings and "not valid JSON" in result.warnings[0]
    assert isinstance(result, Extracted)


def test_zero_byte_files_come_back_as_warnings():
    for name in ("empty.csv", "empty.xlsx", "empty.json", "empty.txt", "empty.pdf", "empty.docx"):
        result = extract(name, b"")
        assert result.warnings, name
        assert result.tables is None, name
        assert isinstance(result.profile, dict), name
    assert extract("empty.csv", b"").profile == {"sheets": [], "primary_sheet": None}
    assert extract("empty.png", b"").profile == {}


def test_an_xlsx_that_is_not_really_a_zip_is_a_warning():
    result = extract("fake.xlsx", b"id,name\n1,a\n")
    assert result.kind == KIND_TABLE
    assert result.tables is None
    assert result.profile == {"sheets": [], "primary_sheet": None}
    assert result.warnings and "workbook" in result.warnings[0].lower()


def test_docx_text_and_heading_extraction():
    result = extract("review.docx", _docx_bytes())

    assert result.kind == KIND_DOCUMENT
    assert result.profile["headings"] == ["Quarterly Review", "Risks"]
    assert "Revenue grew in every region." in result.text
    # Tables are flattened to tab-separated lines in reading order.
    assert "Region\tRevenue" in result.text
    assert "EMEA\t1200" in result.text
    assert result.text.index("Risks") < result.text.index("Region\tRevenue")
    assert result.profile["pages"] is None
    assert result.profile["words"] > 5
    assert result.profile["truncated"] is False


def test_text_file_records_counts_and_markdown_headings():
    body = b"# Title\n\nsome body text here\n\n## Section\n\nmore text\n"
    result = extract("notes.md", body)
    assert result.kind == KIND_DOCUMENT
    assert result.profile["headings"] == ["Title", "Section"]
    assert result.profile["chars"] == len(result.text)
    assert result.profile["words"] == len(result.text.split())
    assert "Title" in result.profile["preview"]


def test_pdf_pages_are_labelled_for_citation():
    data = _pdf_bytes()
    if data is None:
        raise RuntimeError("no PDF writer available")
    result = extract("report.pdf", data)
    assert result.kind == KIND_DOCUMENT
    assert result.profile["pages"] == 2, result.profile
    assert "[page 1]" in result.text and "[page 2]" in result.text
    assert "Alpha" in result.text
    assert result.text.index("[page 1]") < result.text.index("[page 2]")


def test_image_carries_no_extraction():
    result = extract("chart.png", b"\x89PNG\r\n\x1a\n" + b"0" * 100)
    assert result.kind == KIND_IMAGE
    assert result.profile == {}
    assert result.tables is None
    assert result.text == ""
    assert result.warnings == []


def test_row_cap_marks_the_sheet_truncated():
    import services.file_extract as file_extract

    original = file_extract.MAX_ROWS
    file_extract.MAX_ROWS = 5
    try:
        csv_bytes = b"n\n" + b"".join(f"{i}\n".encode() for i in range(20))
        result = extract("many.csv", csv_bytes)
    finally:
        file_extract.MAX_ROWS = original

    sheet = result.profile["sheets"][0]
    assert sheet["rows"] == 5
    assert sheet["truncated"] is True
    assert any("rows" in w for w in result.warnings), result.warnings
    assert len(load_table(result.tables)) == 5


def test_profile_is_json_serializable_with_no_nan_or_numpy():
    import pandas as pd

    frame = pd.DataFrame({"n": [1, None, 3], "s": ["a", None, "c"], "f": [1.5, 2.5, None]})
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        frame.to_excel(writer, sheet_name="Sheet1", index=False)
    result = extract("nulls.xlsx", buffer.getvalue())

    encoded = json.dumps(result.profile)
    assert "NaN" not in encoded and "Infinity" not in encoded
    columns = {c["name"]: c for c in result.profile["sheets"][0]["columns"]}
    assert columns["n"]["nulls"] == 1
    for column in columns.values():
        for value in column["sample_values"]:
            assert value is None or isinstance(value, (int, float, str, bool)), value
            assert type(value).__module__ == "builtins", value


def test_file_card_stays_small_for_a_wide_table():
    import pandas as pd

    frame = pd.DataFrame({f"column_name_number_{i:03d}": [f"value {i}"] * 3 for i in range(200)})
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        frame.to_excel(writer, sheet_name="Wide", index=False)
    result = extract("wide.xlsx", buffer.getvalue())

    card = file_card("wide.xlsx", result.kind, 26_214_400, result.profile)
    assert len(card) < 1500, len(card)
    assert "wide.xlsx" in card
    assert "25.0 MB" in card
    assert "more columns" in card
    assert "3 rows" in card


def test_file_card_shows_sheets_and_sample_rows_for_a_narrow_table():
    data = _workbook_bytes(
        {
            "Summary": [{"metric": "revenue", "value": 10}],
            "Detail": [{"sku": "S1", "qty": 2}, {"sku": "S2", "qty": 3}],
        }
    )
    result = extract("book.xlsx", data)
    card = file_card("book.xlsx", result.kind, 4096, result.profile)

    assert len(card) < 1500
    assert "Summary" in card and "Detail" in card
    assert "Sheet `Detail`" in card
    assert "Sample rows:" in card
    assert "| sku | qty |" in card
    assert "S1" in card


def test_file_card_for_a_document_and_for_json():
    doc = extract("notes.md", b"# Title\n\n" + b"word " * 400)
    card = file_card("notes.md", doc.kind, 2048, doc.profile)
    assert len(card) < 1500
    assert "Headings: Title" in card
    assert "words" in card

    data = extract("config.json", json.dumps({f"key{i}": i for i in range(40)}).encode())
    card = file_card("config.json", data.kind, 900, data.profile)
    assert len(card) < 1500
    assert "JSON object" in card
    assert "… 30 more" in card


def test_file_card_survives_an_empty_profile():
    for kind in (KIND_TABLE, KIND_DOCUMENT, KIND_DATA, KIND_IMAGE):
        card = file_card("mystery.bin", kind, 0, {})
        assert card.startswith("**mystery.bin**"), card
        assert len(card) < 1500


if __name__ == "__main__":
    tests = [
        test_supported_extensions_map_to_the_four_kinds,
        test_sniff_kind_uses_extension_then_mime,
        test_sniff_kind_rejects_an_unsupported_extension,
        test_multi_sheet_workbook_round_trips_by_sheet_name,
        test_csv_with_a_non_utf8_encoding_and_semicolon_delimiter,
        test_tsv_yields_a_single_sheet_named_data,
        test_json_array_of_objects_becomes_a_table,
        test_ndjson_of_objects_becomes_a_table,
        test_json_object_stays_data_with_keys_and_depth,
        test_invalid_json_warns_instead_of_raising,
        test_zero_byte_files_come_back_as_warnings,
        test_an_xlsx_that_is_not_really_a_zip_is_a_warning,
        test_docx_text_and_heading_extraction,
        test_text_file_records_counts_and_markdown_headings,
        test_pdf_pages_are_labelled_for_citation,
        test_image_carries_no_extraction,
        test_row_cap_marks_the_sheet_truncated,
        test_profile_is_json_serializable_with_no_nan_or_numpy,
        test_file_card_stays_small_for_a_wide_table,
        test_file_card_shows_sheets_and_sample_rows_for_a_narrow_table,
        test_file_card_for_a_document_and_for_json,
        test_file_card_survives_an_empty_profile,
    ]
    for test in tests:
        if test is test_pdf_pages_are_labelled_for_citation and _pdf_bytes() is None:
            print(f"SKIP {test.__name__} (no PDF writer installed)")
            continue
        test()
        print(f"PASS {test.__name__}")
