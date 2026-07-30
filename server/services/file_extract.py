"""Parsing user-uploaded chat attachments into something an agent can reason about.

An attachment arrives as up to 25 MB of opaque bytes, and the model that has to
answer questions about it can afford roughly a thousand characters of context.
That gap is the whole reason this module exists: it splits every upload into a
*constant-size* summary the caller can paste into a system prompt
(`file_card`) and a *queryable* payload the caller can hand to a tool later
(`Extracted.tables`, `Extracted.text`). The model never sees the bytes.

Non-obvious decisions:

- **Parquet members inside one zip, not one Parquet file and not extra rows in
  Postgres.** A workbook is several unrelated frames with different columns, so
  a single Parquet file would mean either a union schema or a synthetic
  `sheet` column that pollutes every query. Storing each sheet as its own
  Parquet member keeps dtypes exact and lets `load_table` read one sheet
  without materializing the rest, while the caller still gets *one* `bytea` to
  write — the alternative, a second table keyed by sheet, made saving an
  upload non-atomic. The zip also carries a `manifest.json` member mapping real
  sheet names to sanitized member names, so `load_table` / `sheet_names` work
  from the archive alone and never need the profile to have been persisted
  intact alongside it.
- **Row and character caps.** 1,000,000 rows and 5,000,000 characters are not
  about the 25 MB upload limit; they bound what a *compressed* file can expand
  into. A 3 MB CSV of repeated short rows decompresses into enough DataFrame to
  push a shared app worker into the OOM killer, and every worker shares one
  container. Truncation is recorded in the profile so a later answer can say the
  numbers cover only part of the file rather than silently under-reporting.
- **Corruption is data, not an error.** Users attach half-downloaded PDFs and
  `.xlsx` files that are really CSVs. `extract` reports those as warnings with
  an empty-but-valid profile, because a traceback in the upload path costs the
  user their file and tells them nothing. `UnsupportedFile` is reserved for
  types this module genuinely does not handle, which is the one case the caller
  can act on by rejecting the upload up front.

Heavy third-party imports happen inside the functions that need them: the
routes layer imports this module to call `sniff_kind` on every request, and
paying for pandas and pypdf there would slow app startup for nothing.
"""
from __future__ import annotations

import csv
import io
import json
import math
import os
import re
import zipfile
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

KIND_TABLE = "table"
KIND_DOCUMENT = "document"
KIND_IMAGE = "image"
KIND_DATA = "data"

MAX_ROWS = 1_000_000
MAX_CHARS = 5_000_000

# Preview budgets. The document preview lands in the card verbatim; the JSON
# preview is also exposed to tools, so it gets more room.
PREVIEW_CHARS = 600
JSON_PREVIEW_CHARS = 2_000

CARD_BUDGET = 1_500

# Members of the tables archive. The manifest makes the archive self-describing
# so load_table() does not depend on the caller round-tripping the profile.
MANIFEST_MEMBER = "manifest.json"

_EXTENSIONS: Dict[str, str] = {
    ".csv": KIND_TABLE,
    ".tsv": KIND_TABLE,
    ".xlsx": KIND_TABLE,
    ".xlsm": KIND_TABLE,
    ".pdf": KIND_DOCUMENT,
    ".docx": KIND_DOCUMENT,
    ".txt": KIND_DOCUMENT,
    ".md": KIND_DOCUMENT,
    ".png": KIND_IMAGE,
    ".jpg": KIND_IMAGE,
    ".jpeg": KIND_IMAGE,
    ".webp": KIND_IMAGE,
    ".gif": KIND_IMAGE,
    ".json": KIND_DATA,
    ".ndjson": KIND_DATA,
}

# Browsers disagree with each other about spreadsheet and markdown types, and
# some send nothing at all, so MIME is only consulted when the extension is
# missing or unknown.
_MIME_KINDS: Dict[str, str] = {
    "text/csv": KIND_TABLE,
    "application/csv": KIND_TABLE,
    "text/tab-separated-values": KIND_TABLE,
    "application/vnd.ms-excel": KIND_TABLE,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": KIND_TABLE,
    "application/vnd.ms-excel.sheet.macroenabled.12": KIND_TABLE,
    "application/pdf": KIND_DOCUMENT,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": KIND_DOCUMENT,
    "text/plain": KIND_DOCUMENT,
    "text/markdown": KIND_DOCUMENT,
    "text/x-markdown": KIND_DOCUMENT,
    "application/json": KIND_DATA,
    "text/json": KIND_DATA,
    "application/x-ndjson": KIND_DATA,
    "application/ndjson": KIND_DATA,
}

_UNSAFE_MEMBER_CHARS = re.compile(r"[^A-Za-z0-9._-]+")
_MD_HEADING = re.compile(r"^\s{0,3}(#{1,6})\s+(.+?)\s*#*\s*$")
_BLANK_RUN = re.compile(r"\n{3,}")
_WHITESPACE_RUN = re.compile(r"\s+")


class UnsupportedFile(Exception):
    """The extension/MIME pair is not something this module claims to parse."""


@dataclass
class Extracted:
    kind: str
    profile: Dict[str, Any] = field(default_factory=dict)
    tables: Optional[bytes] = None
    text: str = ""
    warnings: List[str] = field(default_factory=list)


def supported_extensions() -> Dict[str, str]:
    """Extension (with dot, lowercase) to the KIND_* it maps to."""
    return dict(_EXTENSIONS)


def sniff_kind(filename: str, mime: str = "") -> str:
    """Classify an upload, raising UnsupportedFile when we cannot handle it."""
    ext = _extension(filename)
    if ext in _EXTENSIONS:
        return _EXTENSIONS[ext]

    normalized = (mime or "").split(";")[0].strip().lower()
    if normalized in _MIME_KINDS:
        return _MIME_KINDS[normalized]
    # Image subtypes multiply faster than the list above; any image is handled
    # identically anyway (we extract nothing from it).
    if normalized.startswith("image/"):
        return KIND_IMAGE

    label = ext or normalized or "unknown"
    raise UnsupportedFile(
        f"Cannot read '{os.path.basename(filename or 'file')}' ({label}). "
        f"Supported: {', '.join(sorted(_EXTENSIONS))}"
    )


def extract(filename: str, data: bytes, mime: str = "") -> Extracted:
    """Parse `data` according to its type. Only UnsupportedFile ever escapes."""
    kind = sniff_kind(filename, mime)
    ext = _extension(filename)
    data = data or b""
    warnings: List[str] = []

    try:
        if kind == KIND_TABLE:
            return _extract_tabular(ext, data, warnings)
        if kind == KIND_DOCUMENT:
            return _extract_document(ext, data, warnings)
        if kind == KIND_DATA:
            return _extract_data(ext, data, warnings)
        return Extracted(kind=KIND_IMAGE, profile={}, tables=None, text="", warnings=warnings)
    except Exception as exc:  # noqa: BLE001 - a bad upload must not 500 the request
        warnings.append(f"Could not fully read this file: {_describe(exc)}")
        return Extracted(kind=kind, profile=_empty_profile(kind), tables=None, text="", warnings=warnings)


# --------------------------------------------------------------------------
# tables
# --------------------------------------------------------------------------

def _extract_tabular(ext: str, data: bytes, warnings: List[str]) -> Extracted:
    if ext in (".xlsx", ".xlsm"):
        frames = _read_workbook(data, warnings)
    else:
        frames = _read_delimited(ext, data, warnings)
    archive, profile = _package_frames(frames, warnings)
    return Extracted(kind=KIND_TABLE, profile=profile, tables=archive, text="", warnings=warnings)


def _read_delimited(ext: str, data: bytes, warnings: List[str]) -> List[Tuple[str, Any, bool]]:
    import pandas as pd

    text = _decode(data, warnings)
    if not text.strip():
        warnings.append("File is empty; there are no rows to read.")
        return []

    sep = "\t" if ext == ".tsv" else _sniff_delimiter(text)
    reader_kwargs = {"sep": sep, "nrows": MAX_ROWS + 1}
    try:
        frame = pd.read_csv(io.StringIO(text), **reader_kwargs)
    except Exception as exc:  # noqa: BLE001
        try:
            # Ragged rows are common in hand-edited exports and are worth
            # salvaging; anything else is genuinely unreadable.
            frame = pd.read_csv(io.StringIO(text), on_bad_lines="skip", **reader_kwargs)
            warnings.append("Some rows had the wrong number of fields and were skipped.")
        except Exception:  # noqa: BLE001
            warnings.append(f"Could not parse delimited text: {_describe(exc)}")
            return []

    frame, truncated = _cap_rows(frame, "data", warnings)
    return [("data", frame, truncated)]


def _read_workbook(data: bytes, warnings: List[str]) -> List[Tuple[str, Any, bool]]:
    import pandas as pd

    if not data:
        warnings.append("File is empty; there are no sheets to read.")
        return []
    try:
        sheets = pd.read_excel(io.BytesIO(data), sheet_name=None, engine="openpyxl")
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"Could not open this workbook: {_describe(exc)}")
        return []

    frames: List[Tuple[str, Any, bool]] = []
    for name, frame in (sheets or {}).items():
        label = str(name) if name is not None else "sheet"
        capped, truncated = _cap_rows(frame, label, warnings)
        frames.append((label, capped, truncated))
    if not frames:
        warnings.append("This workbook has no sheets.")
    return frames


def _cap_rows(frame: Any, label: str, warnings: List[str]) -> Tuple[Any, bool]:
    if len(frame) > MAX_ROWS:
        warnings.append(f"Sheet '{label}' has more than {MAX_ROWS:,} rows; only the first {MAX_ROWS:,} were kept.")
        return frame.head(MAX_ROWS), True
    return frame, False


def _package_frames(
    frames: List[Tuple[str, Any, bool]], warnings: List[str]
) -> Tuple[Optional[bytes], Dict[str, Any]]:
    """Serialize every frame into one zip of Parquet members plus its profile."""
    used: set = {MANIFEST_MEMBER}
    payloads: List[Tuple[str, bytes]] = []
    sheets: List[Dict[str, Any]] = []

    for name, frame, truncated in frames:
        prepared = _normalize_columns(frame, name, warnings)
        member = _member_name(name, used)
        payload = _to_parquet(prepared, name, warnings)
        if payload is None:
            continue
        payloads.append((member, payload))
        sheets.append(
            {
                "name": name,
                "member": member,
                "rows": int(len(prepared)),
                "truncated": bool(truncated),
                "columns": _column_profiles(prepared),
            }
        )

    if not sheets:
        return None, _empty_profile(KIND_TABLE)

    primary = max(sheets, key=lambda sheet: sheet["rows"])["name"]
    manifest = {
        "sheets": [{"name": s["name"], "member": s["member"], "rows": s["rows"]} for s in sheets],
        "primary_sheet": primary,
    }

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(MANIFEST_MEMBER, json.dumps(manifest))
        for member, payload in payloads:
            archive.writestr(member, payload)

    return buffer.getvalue(), {"sheets": sheets, "primary_sheet": primary}


def _normalize_columns(frame: Any, label: str, warnings: List[str]) -> Any:
    """Parquet requires unique string column names; spreadsheets supply neither."""
    columns = [str(col) for col in frame.columns]
    if len(set(columns)) != len(columns):
        seen: Dict[str, int] = {}
        deduped = []
        for col in columns:
            if col in seen:
                seen[col] += 1
                deduped.append(f"{col}.{seen[col]}")
            else:
                seen[col] = 0
                deduped.append(col)
        columns = deduped
        warnings.append(f"Sheet '{label}' had duplicate column names; later copies were suffixed.")
    if columns != list(frame.columns):
        frame = frame.copy()
        frame.columns = columns
    return frame


def _to_parquet(frame: Any, label: str, warnings: List[str]) -> Optional[bytes]:
    buffer = io.BytesIO()
    try:
        frame.to_parquet(buffer, index=False, engine="pyarrow")
        return buffer.getvalue()
    except Exception:  # noqa: BLE001
        pass

    # Arrow refuses a column holding several Python types at once, which is what
    # a spreadsheet column of "mostly numbers, sometimes N/A" looks like. Text is
    # lossy but keeps the sheet queryable.
    coerced = frame.copy()
    for col in coerced.columns:
        if coerced[col].dtype == object:
            coerced[col] = coerced[col].map(_stringify_cell)
    buffer = io.BytesIO()
    try:
        coerced.to_parquet(buffer, index=False, engine="pyarrow")
        warnings.append(f"Sheet '{label}' had mixed-type columns; they were stored as text.")
        return buffer.getvalue()
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"Sheet '{label}' could not be stored for querying: {_describe(exc)}")
        return None


def _stringify_cell(value: Any) -> Optional[str]:
    if value is None or _is_missing(value):
        return None
    if isinstance(value, str):
        return value
    if isinstance(value, (dict, list, tuple)):
        try:
            return json.dumps(value, default=str)
        except Exception:  # noqa: BLE001
            return str(value)
    return str(value)


def _column_profiles(frame: Any) -> List[Dict[str, Any]]:
    profiles: List[Dict[str, Any]] = []
    for col in frame.columns:
        series = frame[col]
        try:
            nulls = int(series.isna().sum())
        except Exception:  # noqa: BLE001
            nulls = 0
        try:
            unique = int(series.nunique(dropna=True))
        except Exception:  # noqa: BLE001 - unhashable cells (lists, dicts)
            unique = 0
        try:
            samples = [_json_scalar(v) for v in series.dropna().head(3).tolist()]
        except Exception:  # noqa: BLE001
            samples = []
        profiles.append(
            {
                "name": str(col),
                "dtype": str(series.dtype),
                "nulls": nulls,
                "unique": unique,
                "sample_values": samples,
            }
        )
    return profiles


def _member_name(sheet: str, used: set) -> str:
    stem = _UNSAFE_MEMBER_CHARS.sub("_", str(sheet)).strip("._-")[:60] or "sheet"
    candidate = f"{stem}.parquet"
    suffix = 2
    while candidate in used:
        candidate = f"{stem}_{suffix}.parquet"
        suffix += 1
    used.add(candidate)
    return candidate


def load_table(tables: bytes, sheet: Optional[str] = None) -> "pandas.DataFrame":  # noqa: F821
    """Read one sheet back out of the archive; `None` means the primary sheet."""
    import pandas as pd

    manifest = _read_manifest(tables)
    entries = manifest.get("sheets") or []
    if not entries:
        raise KeyError("This archive contains no sheets.")

    wanted = sheet if sheet is not None else manifest.get("primary_sheet")
    entry = _find_sheet(entries, wanted)
    if entry is None:
        available = ", ".join(str(e.get("name")) for e in entries)
        raise KeyError(f"No sheet named '{sheet}'. Available: {available}")

    with zipfile.ZipFile(io.BytesIO(tables)) as archive:
        payload = archive.read(entry["member"])
    return pd.read_parquet(io.BytesIO(payload), engine="pyarrow")


def sheet_names(tables: bytes) -> List[str]:
    """Sheet names in the order they were written."""
    manifest = _read_manifest(tables)
    return [str(entry.get("name")) for entry in (manifest.get("sheets") or [])]


def _read_manifest(tables: bytes) -> Dict[str, Any]:
    if not tables:
        raise ValueError("This file has no stored tables.")
    with zipfile.ZipFile(io.BytesIO(tables)) as archive:
        return json.loads(archive.read(MANIFEST_MEMBER).decode("utf-8"))


def _find_sheet(entries: List[Dict[str, Any]], wanted: Optional[str]) -> Optional[Dict[str, Any]]:
    if wanted is None:
        return entries[0]
    for entry in entries:
        if str(entry.get("name")) == str(wanted):
            return entry
    # Sheet names reach us through a model or a URL, so casing and stray spaces
    # are worth forgiving before failing.
    target = str(wanted).strip().lower()
    for entry in entries:
        if str(entry.get("name")).strip().lower() == target:
            return entry
    for entry in entries:
        if str(entry.get("member")).lower() == target:
            return entry
    return None


# --------------------------------------------------------------------------
# documents
# --------------------------------------------------------------------------

def _extract_document(ext: str, data: bytes, warnings: List[str]) -> Extracted:
    pages: Optional[int] = None
    headings: List[str] = []

    if ext == ".pdf":
        text, pages = _read_pdf(data, warnings)
    elif ext == ".docx":
        text, headings = _read_docx(data, warnings)
    else:
        text = _decode(data, warnings)
        if ext == ".md":
            headings = _markdown_headings(text)
        if not text.strip():
            warnings.append("File is empty; there is no text to read.")

    truncated = False
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS]
        truncated = True
        warnings.append(f"Text was longer than {MAX_CHARS:,} characters; only the beginning was kept.")

    profile = {
        "pages": pages,
        "chars": len(text),
        "words": len(text.split()),
        # Headings are advertised to the model as a table of contents, so a
        # 900-heading document must not blow up the stored profile.
        "headings": [h for h in headings[:100]],
        "preview": _preview(text, PREVIEW_CHARS),
        "truncated": truncated,
    }
    return Extracted(kind=KIND_DOCUMENT, profile=profile, tables=None, text=text, warnings=warnings)


def _read_pdf(data: bytes, warnings: List[str]) -> Tuple[str, Optional[int]]:
    from pypdf import PdfReader

    if not data:
        warnings.append("File is empty; there is no text to read.")
        return "", None
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"Could not open this PDF: {_describe(exc)}")
        return "", None

    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:  # noqa: BLE001
            warnings.append("This PDF is password protected; no text could be read.")
            return "", None

    try:
        pages = len(reader.pages)
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"Could not count pages: {_describe(exc)}")
        return "", None

    chunks: List[str] = []
    total = 0
    empty_pages = 0
    for index in range(pages):
        try:
            body = (reader.pages[index].extract_text() or "").strip()
        except Exception as exc:  # noqa: BLE001
            body = ""
            warnings.append(f"Page {index + 1} could not be read: {_describe(exc)}")
        if not body:
            empty_pages += 1
        # The page marker is a citation anchor: a later tool answers "where does
        # it say that" by searching for it, so it is emitted even for empty pages.
        chunk = f"[page {index + 1}]\n{body}"
        chunks.append(chunk)
        total += len(chunk)
        if total > MAX_CHARS:
            warnings.append(f"Stopped reading after page {index + 1}; the document exceeded the text limit.")
            break

    if empty_pages == pages and pages:
        warnings.append("No selectable text found; this PDF is probably scanned images.")
    return "\n\n".join(chunks), pages


def _read_docx(data: bytes, warnings: List[str]) -> Tuple[str, List[str]]:
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    if not data:
        warnings.append("File is empty; there is no text to read.")
        return "", []
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"Could not open this document: {_describe(exc)}")
        return "", []

    lines: List[str] = []
    headings: List[str] = []
    # Walking the body elements rather than doc.paragraphs + doc.tables keeps
    # reading order, which is what makes the text usable as a citation source.
    for element in document.element.body.iterchildren():
        tag = element.tag.split("}")[-1]
        try:
            if tag == "p":
                paragraph = Paragraph(element, document)
                body = (paragraph.text or "").strip()
                style = getattr(getattr(paragraph, "style", None), "name", "") or ""
                if body and style.startswith("Heading"):
                    headings.append(body)
                if body:
                    lines.append(body)
            elif tag == "tbl":
                for row in Table(element, document).rows:
                    cells = [(cell.text or "").strip().replace("\t", " ") for cell in row.cells]
                    if any(cells):
                        lines.append("\t".join(cells))
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"Part of this document could not be read: {_describe(exc)}")

    if not lines:
        warnings.append("No text found in this document.")
    return "\n".join(lines), headings


def _markdown_headings(text: str) -> List[str]:
    headings: List[str] = []
    for line in text.splitlines():
        match = _MD_HEADING.match(line)
        if match:
            headings.append(match.group(2).strip())
    return headings


# --------------------------------------------------------------------------
# json / ndjson
# --------------------------------------------------------------------------

def _extract_data(ext: str, data: bytes, warnings: List[str]) -> Extracted:
    text = _decode(data, warnings)
    if not text.strip():
        warnings.append("File is empty; there is no JSON to read.")
        return Extracted(
            kind=KIND_DATA,
            profile={"json_kind": "empty", "keys": [], "depth": 0, "preview": ""},
            tables=None,
            text="",
            warnings=warnings,
        )

    parsed, ok = _parse_json(ext, text, warnings)
    if not ok:
        return Extracted(
            kind=KIND_DATA,
            profile={
                "json_kind": "invalid",
                "keys": [],
                "depth": 0,
                "preview": _clip(text.strip(), JSON_PREVIEW_CHARS),
            },
            tables=None,
            text="",
            warnings=warnings,
        )

    records = _as_records(parsed)
    if records is not None:
        import pandas as pd

        truncated = False
        if len(records) > MAX_ROWS:
            warnings.append(f"Kept the first {MAX_ROWS:,} of {len(records):,} records.")
            records = records[:MAX_ROWS]
            truncated = True
        try:
            frame = pd.DataFrame(records)
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"Could not build a table from these records: {_describe(exc)}")
            frame = None
        if frame is not None:
            archive, profile = _package_frames([("data", frame, truncated)], warnings)
            if archive is not None:
                return Extracted(kind=KIND_TABLE, profile=profile, tables=archive, text="", warnings=warnings)

    try:
        preview = json.dumps(parsed, indent=2, default=str)
    except Exception:  # noqa: BLE001
        preview = str(parsed)
    profile = {
        "json_kind": _json_kind(parsed),
        "keys": [str(k) for k in list(parsed.keys())[:100]] if isinstance(parsed, dict) else [],
        "depth": _depth(parsed),
        "preview": _clip(preview, JSON_PREVIEW_CHARS),
    }
    return Extracted(kind=KIND_DATA, profile=profile, tables=None, text="", warnings=warnings)


def _parse_json(ext: str, text: str, warnings: List[str]) -> Tuple[Any, bool]:
    if ext != ".ndjson":
        try:
            return json.loads(text), True
        except Exception as exc:  # noqa: BLE001
            # A `.json` file holding one object per line is a common export
            # mistake and is trivially recoverable; anything else is invalid.
            lines, failures = _parse_ndjson(text)
            if len(lines) >= 2 and not failures:
                warnings.append("This file is newline-delimited JSON, not a single document; read it as records.")
                return lines, True
            warnings.append(f"This file is not valid JSON: {_describe(exc)}")
            return None, False

    lines, failures = _parse_ndjson(text)
    if failures:
        warnings.append(f"{failures} line(s) were not valid JSON and were skipped.")
    if not lines:
        warnings.append("No valid JSON records were found.")
        return None, False
    return lines, True


def _parse_ndjson(text: str) -> Tuple[List[Any], int]:
    values: List[Any] = []
    failures = 0
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        try:
            values.append(json.loads(stripped))
        except Exception:  # noqa: BLE001
            failures += 1
    return values, failures


def _as_records(value: Any) -> Optional[List[Dict[str, Any]]]:
    if isinstance(value, list) and value and all(isinstance(item, dict) for item in value):
        return value
    return None


def _json_kind(value: Any) -> str:
    if value is None:
        return "null"
    if isinstance(value, bool):
        return "boolean"
    if isinstance(value, dict):
        return "object"
    if isinstance(value, list):
        return "array"
    if isinstance(value, (int, float)):
        return "number"
    if isinstance(value, str):
        return "string"
    return type(value).__name__


def _depth(value: Any, level: int = 0) -> int:
    """Container nesting levels; a scalar is 0 and `{"a": 1}` is 1."""
    if level >= 60:
        return level
    if isinstance(value, dict):
        children = list(value.values())
    elif isinstance(value, list):
        # A long array is uniform in practice, so sampling the head keeps this
        # from walking millions of elements for one profile number.
        children = value[:1000]
    else:
        return level
    return max([_depth(child, level + 1) for child in children] or [level + 1])


# --------------------------------------------------------------------------
# the model-facing card
# --------------------------------------------------------------------------

# Successive attempts at the card, from generous to austere. The first that fits
# the budget wins, so a narrow table keeps its sample rows and a 200-column one
# degrades instead of being chopped mid-row.
_CARD_STEPS: Tuple[Tuple[int, int, int, int], ...] = (
    (12, 3, 40, 600),
    (12, 2, 40, 480),
    (8, 2, 32, 360),
    (6, 2, 24, 240),
    (4, 1, 20, 160),
    (2, 0, 16, 80),
)


def file_card(filename: str, kind: str, size_bytes: int, profile: dict) -> str:
    """Compact markdown for a system prompt. Constant size, whatever the file."""
    profile = profile or {}
    label = _clip(os.path.basename(filename or "file"), 60)
    head = f"**{label}** — {kind}, {_human_size(size_bytes)}"

    card = head
    for max_cols, sample_rows, cell, preview in _CARD_STEPS:
        if kind == KIND_TABLE:
            body = _table_body(profile, max_cols, sample_rows, cell)
        elif kind == KIND_DOCUMENT:
            body = _document_body(profile, preview)
        elif kind == KIND_DATA:
            body = _data_body(profile, preview)
        else:
            body = "No text to extract from an image."
        card = f"{head}\n{body}".strip()
        if len(card) <= CARD_BUDGET - 50:
            return card
    return _clip(card, CARD_BUDGET)


def _table_body(profile: dict, max_cols: int, sample_rows: int, cell: int) -> str:
    sheets = profile.get("sheets") or []
    if not sheets:
        return "No readable sheets."

    lines: List[str] = []
    if len(sheets) > 1:
        shown = [str(s.get("name")) for s in sheets[:8]]
        extra = len(sheets) - len(shown)
        listed = ", ".join(_clip(name, 24) for name in shown)
        lines.append(f"Sheets ({len(sheets)}): {listed}" + (f", … {extra} more" if extra > 0 else ""))

    primary_name = profile.get("primary_sheet")
    primary = _find_sheet(sheets, primary_name) or sheets[0]
    columns = primary.get("columns") or []
    rows = primary.get("rows") or 0
    note = ", truncated" if primary.get("truncated") else ""
    lines.append(f"Sheet `{_clip(str(primary.get('name')), 24)}`: {rows:,} rows × {len(columns)} columns{note}")

    shown_cols = columns[:max_cols]
    lines.append("")
    lines.append("| column | dtype |")
    lines.append("| --- | --- |")
    for column in shown_cols:
        lines.append(f"| {_cell(column.get('name'), cell)} | {_cell(column.get('dtype'), 16)} |")
    if len(columns) > len(shown_cols):
        lines.append(f"… {len(columns) - len(shown_cols)} more columns")

    sample = _sample_rows(shown_cols, sample_rows, cell)
    if sample:
        lines.append("")
        lines.append("Sample rows:")
        lines.extend(sample)
    return "\n".join(lines)


def _sample_rows(columns: List[Dict[str, Any]], limit: int, cell: int) -> List[str]:
    if limit <= 0 or not columns:
        return []
    depth = min(limit, max((len(c.get("sample_values") or []) for c in columns), default=0))
    if depth <= 0:
        return []
    header = "| " + " | ".join(_cell(c.get("name"), cell) for c in columns) + " |"
    divider = "| " + " | ".join("---" for _ in columns) + " |"
    body = []
    for index in range(depth):
        values = []
        for column in columns:
            samples = column.get("sample_values") or []
            values.append(_cell(samples[index], cell) if index < len(samples) else "")
        body.append("| " + " | ".join(values) + " |")
    return [header, divider] + body


def _document_body(profile: dict, preview_chars: int) -> str:
    facts = []
    if profile.get("pages") is not None:
        facts.append(f"{int(profile['pages']):,} pages")
    facts.append(f"{int(profile.get('words') or 0):,} words")
    facts.append(f"{int(profile.get('chars') or 0):,} chars")
    if profile.get("truncated"):
        facts.append("truncated")
    lines = [" · ".join(facts)]

    headings = [str(h) for h in (profile.get("headings") or [])]
    if headings:
        shown = headings[:10]
        listed = "; ".join(_clip(h, 40) for h in shown)
        extra = len(headings) - len(shown)
        lines.append(f"Headings: {listed}" + (f"; … {extra} more" if extra > 0 else ""))

    preview = _flatten(str(profile.get("preview") or ""))
    if preview:
        lines.append(f"Preview: {_clip(preview, preview_chars)}")
    return "\n".join(lines)


def _data_body(profile: dict, preview_chars: int) -> str:
    kind = str(profile.get("json_kind") or "unknown")
    keys = [str(k) for k in (profile.get("keys") or [])]
    lines = [f"JSON {kind} · depth {int(profile.get('depth') or 0)} · {len(keys)} top-level keys"]
    if keys:
        shown = keys[:10]
        extra = len(keys) - len(shown)
        lines.append("Keys: " + ", ".join(_clip(k, 24) for k in shown) + (f", … {extra} more" if extra > 0 else ""))
    preview = _flatten(str(profile.get("preview") or ""))
    if preview:
        lines.append(f"Preview: {_clip(preview, preview_chars)}")
    return "\n".join(lines)


# --------------------------------------------------------------------------
# shared helpers
# --------------------------------------------------------------------------

def _extension(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


def _empty_profile(kind: str) -> Dict[str, Any]:
    if kind == KIND_TABLE:
        return {"sheets": [], "primary_sheet": None}
    if kind == KIND_DOCUMENT:
        return {"pages": None, "chars": 0, "words": 0, "headings": [], "preview": "", "truncated": False}
    if kind == KIND_DATA:
        return {"json_kind": "unknown", "keys": [], "depth": 0, "preview": ""}
    return {}


def _decode(data: bytes, warnings: List[str]) -> str:
    if not data:
        return ""

    candidates: List[str] = []
    try:
        import chardet

        # A sample is enough and bounds the cost on a 25 MB upload.
        guess = chardet.detect(data[:200_000]) or {}
        encoding = (guess.get("encoding") or "").strip()
        if encoding and float(guess.get("confidence") or 0) >= 0.5:
            candidates.append(encoding)
    except Exception:  # noqa: BLE001
        pass

    if data.startswith(b"\xef\xbb\xbf"):
        candidates.append("utf-8-sig")
    candidates.extend(["utf-8", "latin-1"])

    seen: set = set()
    for encoding in candidates:
        key = encoding.lower()
        if key in seen:
            continue
        seen.add(key)
        try:
            return data.decode(encoding)
        except Exception:  # noqa: BLE001
            continue

    warnings.append("Text could not be decoded cleanly; unreadable bytes were replaced.")
    return data.decode("utf-8", errors="replace")


def _sniff_delimiter(text: str) -> str:
    sample = text[:65_536]
    cut = sample.rfind("\n")
    if cut > 0:
        sample = sample[:cut]
    try:
        return csv.Sniffer().sniff(sample, delimiters=",;\t|").delimiter
    except Exception:  # noqa: BLE001
        return ","


def _is_missing(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, float):
        return math.isnan(value)
    try:
        import pandas as pd

        result = pd.isna(value)
        return bool(result) if not hasattr(result, "__len__") else False
    except Exception:  # noqa: BLE001
        return False


def _json_scalar(value: Any) -> Any:
    """Coerce a cell into something json.dumps and Postgres jsonb both accept."""
    if _is_missing(value):
        return None
    unwrap = getattr(value, "item", None)
    if callable(unwrap) and hasattr(value, "dtype"):
        try:
            value = value.item()
        except Exception:  # noqa: BLE001
            pass
    if isinstance(value, bool):
        return value
    if isinstance(value, int):
        return int(value)
    if isinstance(value, float):
        return float(value) if math.isfinite(value) else None
    if isinstance(value, str):
        return _clip(value, 200)
    if isinstance(value, bytes):
        return _clip(value.decode("utf-8", errors="replace"), 200)
    for attribute in ("isoformat",):
        method = getattr(value, attribute, None)
        if callable(method):
            try:
                return str(method())
            except Exception:  # noqa: BLE001
                break
    return _clip(str(value), 200)


def _clip(text: str, limit: int) -> str:
    text = text or ""
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 1)].rstrip() + "…"


def _cell(value: Any, limit: int) -> str:
    if value is None:
        return ""
    rendered = _flatten(str(value)).replace("|", "\\|")
    return _clip(rendered, limit)


def _flatten(text: str) -> str:
    return _WHITESPACE_RUN.sub(" ", text or "").strip()


def _preview(text: str, limit: int) -> str:
    return _clip(_BLANK_RUN.sub("\n\n", (text or "").strip()), limit)


def _human_size(size_bytes: int) -> str:
    try:
        size = float(size_bytes or 0)
    except Exception:  # noqa: BLE001
        return "unknown size"
    for unit in ("B", "KB", "MB", "GB"):
        if size < 1024 or unit == "GB":
            return f"{int(size)} {unit}" if unit == "B" else f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} GB"


def _describe(exc: Exception) -> str:
    message = _flatten(str(exc)) or exc.__class__.__name__
    return _clip(message, 200)
